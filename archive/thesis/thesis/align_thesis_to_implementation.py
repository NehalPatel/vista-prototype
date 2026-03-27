# -*- coding: utf-8 -*-
"""
Align Vista-Thesis.docx with the vista-prototype implementation (thesis alignment plan).
Replaces Chapters 4, 6–10 body text and Chapter 1 from Chapter-01-Revised.md.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document

from alignment_chapters import (
    CH10_BLOCKS,
    CH4_BLOCKS,
    CH6_BLOCKS,
    CH7_BLOCKS,
    CH8_BLOCKS,
    CH9_BLOCKS,
)
from apply_thesis_fixes import (
    delete_paragraph,
    insert_paragraph_after,
    parse_chapter1_md,
    replace_chapter1,
)

ROOT = Path(__file__).resolve().parent
DOC_PATH = ROOT / "Vista-Thesis.docx"
CH1_MD = ROOT / "Chapter-01-Revised.md"


def find_chapter_start(doc: Document, title_prefix: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(title_prefix):
            return i
    raise RuntimeError(f"Could not find chapter starting with {title_prefix!r}")


def find_next_chapter_start(doc: Document, after_index: int, title_prefix: str) -> int:
    for j in range(after_index + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].text.strip().startswith(title_prefix):
            return j
    return len(doc.paragraphs)


def replace_chapter_body(
    doc: Document,
    chapter_title_prefix: str,
    next_chapter_prefix: str,
    blocks: list[tuple[str, str]],
) -> None:
    start = find_chapter_start(doc, chapter_title_prefix)
    next_i = find_next_chapter_start(doc, start, next_chapter_prefix)
    title_para = doc.paragraphs[start]
    for idx in range(next_i - 1, start, -1):
        delete_paragraph(doc.paragraphs[idx])
    anchor = title_para
    for style, text in blocks:
        np = insert_paragraph_after(anchor)
        try:
            np.style = style
        except Exception:
            np.style = "Normal"
        np.add_run(text)
        anchor = np


def set_paragraph_startswith(doc: Document, prefix: str, new_text: str) -> None:
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            p.text = new_text
            return
    raise RuntimeError(f"No paragraph starts with {prefix!r}")


def patch_chapter3_scope(doc: Document) -> None:
    """Align methodology and requirements in Ch 3 with prototype scope (no search API)."""
    start = find_chapter_start(doc, "Chapter 3: Research Methodology")
    end = find_chapter_start(doc, "Chapter 4:")
    replacements: list[tuple[str, str]] = [
        (
            "Development of the visual search engine",
            "Conceptual specification of retrieval over indexed data (executable search service not implemented in the prototype repository)",
        ),
        (
            "Performance testing using benchmark datasets",
            "Operational testing on sample YouTube videos and author-prepared training data (no fixed benchmark corpus in the repository)",
        ),
        (
            "Evaluation using information retrieval metrics",
            "Planned information-retrieval metrics for future labeled studies; processing-stage timings from run_stats where recorded",
        ),
        (
            "Analysis of accuracy and efficiency",
            "Qualitative review of detections and analysis of pipeline timing and resource use",
        ),
        (
            "Accept natural language queries",
            "Target capability for a future search layer; the prototype UI accepts a video URL for processing, not free-text search queries",
        ),
        (
            "Retrieve relevant frames and videos",
            "Persist frame-level evidence in JSON and optionally MongoDB so a future search service can retrieve matches",
        ),
        (
            "Support efficient querying",
            "Support efficient indexing and field shapes suitable for multi-entity queries once a search service is added",
        ),
        (
            "Efficient indexing and low query latency",
            "Efficient indexing on write; query latency applies to future retrieval implementation",
        ),
        (
            "User-friendly interface for querying",
            "User-friendly web interface for video processing, training data management, and inspection of results",
        ),
        (
            "Metadata Generation → Database Storage → Query Processing → Retrieval Results",
            "Metadata Generation → Optional MongoDB Storage → (Future) Query Processing → (Future) Retrieval Results",
        ),
        (
            "Queries are executed",
            "Indexed data would be queried by a future search component (not part of the current prototype)",
        ),
        (
            "Relevant results are retrieved",
            "Relevant evidence is materialized in JSON and optional MongoDB documents for each processed video",
        ),
    ]
    for i in range(start, end):
        t = doc.paragraphs[i].text
        if not t:
            continue
        new_t = t
        for old, new in replacements:
            if old in new_t:
                new_t = new_t.replace(old, new)
        if new_t != t:
            doc.paragraphs[i].text = new_t


def ensure_chapter6_title(doc: Document) -> None:
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Chapter 6:") and "Visual Search Engine" in t:
            p.text = "Chapter 6: Conceptual Design for Retrieval and Ranking"
            return


def main() -> None:
    doc = Document(str(DOC_PATH))

    patch_chapter3_scope(doc)
    ensure_chapter6_title(doc)

    # Replace from back to front so indices stay valid for title search (we re-find each time).
    replace_chapter_body(
        doc,
        "Chapter 10: Conclusion and Future Work",
        "\uFFFF",
        CH10_BLOCKS,
    )
    replace_chapter_body(doc, "Chapter 9: Discussion", "Chapter 10:", CH9_BLOCKS)
    replace_chapter_body(
        doc,
        "Chapter 8: Experimental Evaluation and Results",
        "Chapter 9:",
        CH8_BLOCKS,
    )
    replace_chapter_body(
        doc,
        "Chapter 7: System Implementation and Integration",
        "Chapter 8:",
        CH7_BLOCKS,
    )
    replace_chapter_body(
        doc,
        "Chapter 6: Conceptual Design for Retrieval and Ranking",
        "Chapter 7:",
        CH6_BLOCKS,
    )
    replace_chapter_body(
        doc,
        "Chapter 4: Architecture and Design of the Vista-prototype",
        "Chapter 5:",
        CH4_BLOCKS,
    )

    ch1_blocks = parse_chapter1_md(CH1_MD)
    replace_chapter1(doc, ch1_blocks)

    doc.save(str(DOC_PATH))
    print(f"Updated {DOC_PATH}")


if __name__ == "__main__":
    main()
