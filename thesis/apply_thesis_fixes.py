# -*- coding: utf-8 -*-
"""
Apply URGENT_THESIS_FIXES.md and Chapter-01-Revised.md to Vista-Thesis.docx.
"""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent
DOC_PATH = ROOT / "Vista-Thesis.docx"
BACKUP_PATH = ROOT / "Vista-Thesis.backup.docx"
CH1_MD = ROOT / "Chapter-01-Revised.md"

def delete_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def insert_paragraph_after(anchor: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    return Paragraph(new_p, anchor._parent)


def strip_emojis(text: str) -> str:
    if not text:
        return text
    out = []
    for ch in text:
        o = ord(ch)
        if o == 0xFE0F or o == 0x200D:
            continue
        if 0x1F300 <= o <= 0x1FAFF:
            continue
        if 0x2600 <= o <= 0x27BF and ch not in "°":
            continue
        if ch in "📘✍📊📌👉🔷✅🚀✔📌":
            continue
        out.append(ch)
    return "".join(out)


def clean_instructional_phrases(text: str) -> str:
    if not text:
        return text
    t = text
    t = re.sub(
        r"\s*\(?\s*very important for examiners\s*\)?\.?",
        "",
        t,
        flags=re.IGNORECASE,
    )
    t = re.sub(r"\s*\(?\s*Very Important for PhD\s*\)?\.?", "", t, flags=re.IGNORECASE)
    t = re.sub(r"\s*\(?\s*VERY IMPORTANT\s*\)?", "", t)
    t = re.sub(r"\s*✔\s*$", "", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def parse_chapter1_md(path: Path) -> list[tuple[str, str]]:
    """Return list of (kind, text) where kind is h1|h2|h3|p."""
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks: list[tuple[str, str]] = []
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("### "):
            blocks.append(("h3", stripped[4:].strip()))
            i += 1
            continue
        if stripped.startswith("## "):
            blocks.append(("h2", stripped[3:].strip()))
            i += 1
            continue
        if stripped.startswith("# ") and not stripped.startswith("##"):
            blocks.append(("h1", stripped[2:].strip()))
            i += 1
            continue
        if re.match(r"^\d+\.\s", stripped):
            blocks.append(("p", stripped))
            i += 1
            continue
        buf = [stripped]
        i += 1
        while i < len(lines):
            L = lines[i].strip()
            if not L:
                break
            if L.startswith("#") or re.match(r"^\d+\.\s", L):
                break
            buf.append(L)
            i += 1
        blocks.append(("p", " ".join(buf)))
    return blocks


def find_chapter1_span(doc: Document) -> tuple[int, int]:
    """Return (ch1_start, ch2_index) so Chapter 1 is paragraphs[ch1_start:ch2_index]."""
    ch2_idx = None
    for i, p in enumerate(doc.paragraphs):
        ts = p.text.strip()
        if ts.startswith("Chapter 2:") and "Literature" in ts:
            ch2_idx = i
            break
    if ch2_idx is None:
        raise RuntimeError("Could not find start of Chapter 2 (Literature Review).")
    ch1_start = None
    for i in range(ch2_idx - 1, -1, -1):
        p = doc.paragraphs[i]
        ts = p.text.strip()
        if (p.style.name == "Heading 1" and "Chapter 1" in ts) or ts.startswith(
            "Chapter 1:"
        ):
            ch1_start = i
            break
    if ch1_start is None:
        ch1_start = 1
    return ch1_start, ch2_idx


def replace_chapter1(doc: Document, blocks: list[tuple[str, str]]) -> None:
    ch1_start, ch2_idx = find_chapter1_span(doc)
    anchor = doc.paragraphs[ch1_start - 1]
    for idx in range(ch2_idx - 1, ch1_start - 1, -1):
        delete_paragraph(doc.paragraphs[idx])
    style_map = {"h1": "Heading 1", "h2": "Heading 2", "h3": "Heading 3", "p": "Normal"}
    for kind, text in blocks:
        np = insert_paragraph_after(anchor)
        try:
            np.style = style_map[kind]
        except KeyError:
            np.style = "Normal"
        np.add_run(text)
        anchor = np


def should_split_merged_lines(first: str, second: str) -> bool:
    if not first or not second:
        return False
    if first.startswith("Query:"):
        return True
    if re.match(r"^Chapter \d+:", first) and second[0].isupper():
        return True
    if len(first) < 100 and not first.endswith(".") and len(second) > 15:
        if first[0].isupper():
            return True
    return False


def split_merged_newlines(doc: Document) -> None:
    """Split paragraphs that cram a title and body into one w:p with a newline."""
    for p in list(doc.paragraphs):
        t = p.text
        if "\n" not in t:
            continue
        parts = t.split("\n")
        if len(parts) != 2:
            continue
        a, b = parts[0].strip(), parts[1].strip()
        if not should_split_merged_lines(a, b):
            continue
        p.text = a
        np = insert_paragraph_after(p)
        np.style = p.style
        np.add_run(b)


def paragraph_is_checklist_section_start(t: str) -> bool:
    s = t.strip()
    return (
        s.startswith("📊 What You Should Add")
        or s.startswith("📌 What You Should Add")
        or s.startswith("📌 What You MUST Add")
    )


def delete_checklist_sections(doc: Document) -> None:
    i = 0
    while i < len(doc.paragraphs):
        t = doc.paragraphs[i].text
        if paragraph_is_checklist_section_start(t):
            while i < len(doc.paragraphs):
                cur = doc.paragraphs[i].text.strip()
                if cur.startswith("📘 Chapter"):
                    break
                delete_paragraph(doc.paragraphs[i])
            continue
        i += 1


def delete_final_suggestions(doc: Document) -> None:
    i = 0
    while i < len(doc.paragraphs):
        if doc.paragraphs[i].text.strip().startswith("🚀 Final Suggestions"):
            while i < len(doc.paragraphs):
                delete_paragraph(doc.paragraphs[i])
            return
        i += 1


def delete_standalone_template_paragraph(t: str) -> bool:
    s = t.strip()
    if not s:
        return False
    if s.startswith("✍️") or s.startswith("Write This"):
        return True
    if s.startswith("👉 Insert"):
        return True
    if s.startswith("🔷 Diagram Name:"):
        return True
    if "Write This" in s and len(s) < 40:
        return True
    if s.startswith("📊 Diagram You MUST Add"):
        return True
    if s.startswith("📊 Optional Mini Table"):
        return True
    if s.startswith("📌 Optional Screenshot"):
        return True
    if s.startswith("📊 Diagram to Add"):
        return True
    if s == "📊 Diagram (VERY IMPORTANT)":
        return True
    if re.match(r"^🔷 Figure \d+\.\d+:", s) and len(s) < 120:
        return True
    if s.startswith("✅") and any(
        k in s
        for k in (
            "MUST ADD",
            "ADD LATER",
            "Add Later",
            "Add References",
            "Add Graphs",
            "Add Real Results",
            "Add Figures",
            "Add Screenshots",
            "Formatting",
            "Diagrams",
            "Tables",
            "Replace",
            "MANDATORY",
            "strong viva",
        )
    ):
        return True
    if s.startswith("You should cite:"):
        return True
    if s in (
        "YOLO papers",
        "ArcFace / InsightFace",
        "CBVR research papers",
        "Before submission, make sure to:",
    ):
        return True
    if s.startswith("Font:") or s.startswith("Size:") or s.startswith("Line spacing:"):
        return True
    if s in ("UI", "Results", "Annotated frames", "System architecture", "Pipeline diagrams", "Graphs (Precision/Recall)"):
        return True
    return False


def remove_placeholder_paragraphs(doc: Document) -> None:
    changed = True
    while changed:
        changed = False
        for p in list(doc.paragraphs):
            if delete_standalone_template_paragraph(p.text):
                delete_paragraph(p)
                changed = True
                break


def polish_all_paragraphs(doc: Document) -> None:
    for p in doc.paragraphs:
        t = p.text
        if not t:
            continue
        t2 = strip_emojis(t)
        t2 = clean_instructional_phrases(t2)
        if t2 != t:
            p.text = t2


def main() -> None:
    if not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)
    doc = Document(str(DOC_PATH))
    blocks = parse_chapter1_md(CH1_MD)
    replace_chapter1(doc, blocks)
    delete_checklist_sections(doc)
    delete_final_suggestions(doc)
    remove_placeholder_paragraphs(doc)
    split_merged_newlines(doc)
    polish_all_paragraphs(doc)
    doc.save(str(DOC_PATH))
    print(f"Saved {DOC_PATH} (backup: {BACKUP_PATH})")


if __name__ == "__main__":
    main()
